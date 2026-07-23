"""
HSE Intelligence Platform — Full Data Flow Word (.docx) Generator
Mirrors generate_pdf.py content, EXCLUDING all contractor-specific material.
"""

import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ───────────────────────────────────────────────────────────────────
TEAL, DARK_BLUE, MID_BLUE = "12B8A6", "1E293B", "334155"
ACCENT, WARN, DANGER, SUCCESS, PURPLE = "0EA5E9", "F59E0B", "EF4444", "22C55E", "8B5CF6"
LIGHT_BG, LIGHT_GREY, TEXT_DARK, TEXT_MID = "F0FDFA", "F1F5F9", "0F172A", "475569"
WHITE = "FFFFFF"

FULL_W = 174  # usable page width in mm (A4 minus 18mm margins each side)

OUTPUT = os.path.join(os.path.dirname(__file__), "HSE_Platform_Data_Flow.docx")

doc = Document()
# A4 + margins
sec = doc.sections[0]
sec.page_width, sec.page_height = Mm(210), Mm(297)
sec.left_margin = sec.right_margin = Mm(18)
sec.top_margin = sec.bottom_margin = Mm(16)
# Default font
normal = doc.styles["Normal"].font
normal.name = "Arial"
normal.size = Pt(9)


# ── Low-level helpers ─────────────────────────────────────────────────────────
def shade(cell, hexfill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hexfill)
    tcPr.append(shd)


def fill_cell(cell, text, bold=False, color=None, size=9, align=None, mono=False, fill=None):
    if fill:
        shade(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    if align is not None:
        p.alignment = align
    lines = str(text).split("\n")
    for li, line in enumerate(lines):
        run = p.add_run(line)
        run.font.size = Pt(size); run.font.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
        if mono:
            run.font.name = "Consolas"
        if li < len(lines) - 1:
            run.add_break()
    return p


def set_widths(table, widths_mm):
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for i, w in enumerate(widths_mm):
            row.cells[i].width = Mm(w)


def spacer(pts=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.space_before = Pt(0)
    return p


# ── Content helpers ───────────────────────────────────────────────────────────
def section_header(title, fill=TEAL):
    t = doc.add_table(rows=1, cols=1)
    set_widths(t, [FULL_W])
    fill_cell(t.cell(0, 0), "  " + title, bold=True, color=WHITE, size=12, fill=fill)
    spacer(3)


def body(text, bold=False, size=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(TEXT_DARK)
    return p


def subhead(text, color=DARK_BLUE, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.bold = True; run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    return p


def note(text):
    t = doc.add_table(rows=1, cols=1)
    set_widths(t, [FULL_W])
    fill_cell(t.cell(0, 0), text, size=8, color="92400E", fill="FEF3C7")
    spacer(3)


def kv_table(rows, mono_keys=("Formula / Logic",)):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = "Table Grid"
    set_widths(t, [50, FULL_W - 50])
    for i, (k, v) in enumerate(rows):
        fill_cell(t.cell(i, 0), k, bold=True, size=8, fill=LIGHT_GREY)
        fill_cell(t.cell(i, 1), v, size=8, mono=(k in mono_keys))
    spacer(2)
    return t


def chart_block(title, what, why, formula, db_source, notes=None, fill=TEAL):
    tt = doc.add_table(rows=1, cols=1)
    set_widths(tt, [FULL_W])
    fill_cell(tt.cell(0, 0), "  " + title, bold=True, color=WHITE, size=10, fill=fill)
    rows = [
        ("What is it?", what),
        ("Why shown?", why),
        ("Formula / Logic", formula),
        ("DB Source", db_source),
    ]
    if notes:
        rows.append(("Notes", notes))
    kv_table(rows)


def grid_table(data, widths, header_fill=DARK_BLUE, first_col_bold=False):
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.style = "Table Grid"
    set_widths(t, widths)
    for r, row in enumerate(data):
        for c, val in enumerate(row):
            is_header = (r == 0)
            fill_cell(
                t.cell(r, c), val,
                bold=is_header or (first_col_bold and c == 0),
                color=WHITE if is_header else None,
                size=8 if is_header else 7.5,
                fill=header_fill if is_header else (LIGHT_BG if r % 2 else WHITE),
            )
    spacer(3)
    return t


# ══════════════════════════════════════════════════════════════════════════════
#  COVER
# ══════════════════════════════════════════════════════════════════════════════
title_p = doc.add_paragraph(); title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("HSE Intelligence Platform")
r.font.size = Pt(26); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(TEAL)
sub_p = doc.add_paragraph(); sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run("Complete Data Flow Reference")
r.font.size = Pt(15); r.font.bold = True; r.font.color.rgb = RGBColor.from_string(MID_BLUE)
tag_p = doc.add_paragraph(); tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = tag_p.add_run("Charts · Formulas · DB Sources · Business Logic")
r.font.size = Pt(9); r.font.color.rgb = RGBColor.from_string("64748B")
spacer(6)

meta = [
    ["Platform", "HSE Intelligence Platform (hse_old_ui)"],
    ["Frontend", "React 18 + TypeScript + Vite + Tailwind CSS v4"],
    ["Backend", "FastAPI (Python) + SQLAlchemy 2.0"],
    ["Database", "MySQL via XAMPP"],
    ["Auth", "JWT (HS256, 60-min TTL) — org_id embedded in token"],
    ["API Base", "http://localhost:8080/api/v1  (proxied from Vite /api/v1)"],
    ["Multi-tenant", "Every table has organisation_id; all queries filter by org"],
    ["Document scope", "Every page, every chart, every KPI — real data sources documented"],
]
kv_table([(k, v) for k, v in meta], mono_keys=())
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════
subhead("Table of Contents", size=16)
toc = [
    ["#", "Page", "What it covers"],
    ["1", "Architecture Overview", "How data flows from DB → API → UI"],
    ["2", "Dashboard  /", "KPI cards, incident trend, heatmap"],
    ["3", "Checklists  /checklists", "Status breakdown, compliance rate, category rings"],
    ["4", "Users  /users", "Employee roster, department stats"],
    ["5", "Root Cause Analysis  /root-cause-analysis", "RCA list, priority, status"],
    ["6", "Actions (CAPA)  /actions", "CAPA table, status KPIs, overdue tracking"],
    ["7", "Compliance  /compliance", "Compliance score, audit readiness, trend"],
    ["8", "Analytics  /analytics", "7 tabs: Violations, Near Miss, RCA, Permits, Zone Risk, Trend, PPE"],
    ["9", "Risk  /risk", "Risk matrix, zone bar chart, residual trend, aging"],
    ["10", "Equipment Certification  /equipment-certification", "Cert status, type breakdown"],
    ["11", "Violations  /violations + /violations/:id", "Incident list, detail page"],
    ["12", "Engagement  /engagement", "Reporting rate, survey score, ring KPIs, CAPA actions"],
    ["13", "Formula Cheat-Sheet", "All computed metrics in one place"],
    ["14", "DB Table Reference", "Table name → columns → which pages use it"],
]
grid_table(toc, [10, 70, FULL_W - 80], first_col_bold=True)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1 — ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
section_header("1.  Architecture Overview")
body("Every data point visible in the UI originates from the MySQL database and passes through "
     "the FastAPI backend before reaching the React frontend. No hard-coded business data exists "
     "anywhere in the final production build.")
arch = [
    ["Layer", "Technology", "Role"],
    ["Database", "MySQL 8 (XAMPP)", "Stores all entities: incidents, employees, safety walks, permits, CAPA, etc."],
    ["ORM", "SQLAlchemy 2.0", "Maps Python classes to DB tables; enforces multi-tenant org filtering"],
    ["Backend", "FastAPI (Python 3.11)", "REST API on port 8080; JWT auth; business logic & aggregations"],
    ["Proxy", "Vite Dev Server", "Forwards /api/* from port 5173 → port 8080 so CORS is bypassed"],
    ["Frontend", "React 18 + TypeScript", "Calls axiosInstance which adds Bearer token from localStorage"],
    ["State", "React useState / useEffect", "Each page fetches its own data on mount; no global cache"],
    ["Auth", "JWT HS256 (60-min TTL)", "org_id + user id embedded; backend extracts with get_current_user()"],
]
grid_table(arch, [30, 42, FULL_W - 72])
subhead("Multi-tenant filtering rule", color=MID_BLUE, size=10)
body("Every SQL query in the backend passes through _org_filter(query, Model, org_id) which appends "
     "WHERE model.organisation_id = :org_id. This means every user only ever sees data belonging to "
     "their organisation.")
code_t = doc.add_table(rows=1, cols=1); set_widths(code_t, [FULL_W])
fill_cell(code_t.cell(0, 0),
          "def _org_filter(query, model, org_id):\n"
          "    if org_id is not None:\n"
          "        return query.filter(model.organisation_id == org_id)\n"
          "    return query",
          size=8, mono=True, fill="EFF6FF", color="1D4ED8")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2 — DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
section_header("2.  Dashboard  ( / )")
body("The dashboard is the landing page. It calls /api/v1/analytics/dashboard-summary and renders "
     "KPI cards, a monthly incident trend chart, a severity pie chart, and a site-activity heatmap.")
chart_block("Total Incidents (KPI Card)",
    "Count of all incident records for the org this month.",
    "Gives managers a quick pulse on incident volume for the current period.",
    "SELECT COUNT(*) FROM incidents WHERE organisation_id=:org AND report_date >= first_of_month",
    "Table: incidents  Columns: id, organisation_id, report_date", fill=DANGER)
chart_block("Near Misses (KPI Card)",
    "Count of near-miss events reported this month.",
    "Near misses are leading indicators of accidents; tracking them drives preventive action.",
    "SELECT COUNT(*) FROM near_misses WHERE organisation_id=:org AND DATE(event_date_time) >= first_of_month",
    "Table: near_misses  Columns: id, organisation_id, event_date_time", fill=WARN)
chart_block("Open CAPA Actions (KPI Card)",
    "Count of corrective/preventive actions that are not yet 'Completed'.",
    "Unresolved CAPA actions represent outstanding safety obligations.",
    "SELECT COUNT(*) FROM capa_actions WHERE organisation_id=:org AND status != 'Completed'",
    "Table: capa_actions  Columns: id, organisation_id, status", fill=PURPLE)
chart_block("Monthly Incident Trend (Line Chart)",
    "Bar or line chart showing incident count per month for the last 10 months.",
    "Reveals whether incident rates are improving or worsening over time.",
    "For each of last 10 months: SELECT COUNT(*) FROM incidents WHERE YEAR=y AND MONTH=m AND org=:org  "
    "-> list of {month: 'JAN', value: N}",
    "Table: incidents  Columns: id, organisation_id, report_date", fill=ACCENT)
chart_block("Severity Distribution (Pie / Donut Chart)",
    "Breakdown of incidents by severity level: Critical / High / Medium / Low.",
    "Identifies whether high-severity incidents dominate, guiding resource prioritisation.",
    "SELECT severity, COUNT(*) FROM incidents WHERE org=:org GROUP BY severity",
    "Table: incidents  Columns: severity (enum string)  Mapping: 1->Critical, 2->High, 3->Medium, 4->Low, else->Unknown",
    fill=DANGER)
chart_block("Site Activity Heatmap",
    "Grid of sites x months showing incident count per cell.",
    "Highlights which sites have persistent high incident rates month-on-month.",
    "For each site s and month m: SELECT COUNT(*) FROM incidents i JOIN working_stations ws "
    "ON i.location_station_id=ws.id WHERE ws.site_id=s AND org=:org AND YEAR/MONTH=m",
    "Tables: incidents, working_stations, sites  Columns: location_station_id, site_id, site.name", fill=MID_BLUE)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3 — CHECKLISTS
# ══════════════════════════════════════════════════════════════════════════════
section_header("3.  Checklists  ( /checklists )")
body("Fetches safety-walk inspection records. API endpoint: GET /api/v1/safety-walks/. "
     "The frontend maps raw DB rows to checklist-style display items.")
chart_block("Status Breakdown (Bar / Donut)",
    "Count of safety walks grouped by follow_up_required (Yes/No) and compliance_rating bands.",
    "Measures how many inspections identified issues requiring follow-up.",
    "SELECT follow_up_required, COUNT(*) FROM safety_walks WHERE org=:org GROUP BY follow_up_required",
    "Table: safety_walks  Columns: id, organisation_id, follow_up_required, compliance_rating, inspection_date_time",
    fill=TEAL)
chart_block("Compliance Rate (KPI %)",
    "Percentage of walks with compliance_rating >= 4 (out of 5).",
    "A high rate means inspections are passing; a low rate signals systemic issues.",
    "compliant_walks / total_walks x 100  where compliant_walks: rating >= 4",
    "Table: safety_walks  Column: compliance_rating (Integer 1-5)", fill=SUCCESS)
chart_block("Inspection Type Rings (Donut Charts)",
    "Each ring shows walks of a specific inspection_type (e.g., General, Toolbox) as a % of total.",
    "Breaks down where safety attention is being directed.",
    "SELECT inspection_type, COUNT(*) FROM safety_walks WHERE org=:org GROUP BY inspection_type",
    "Table: safety_walks  Column: inspection_type (VARCHAR 100)", fill=ACCENT)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4 — USERS
# ══════════════════════════════════════════════════════════════════════════════
section_header("4.  Users  ( /users )")
body("Displays the employee roster. API: GET /api/v1/employees/?limit=200 plus "
     "GET /api/v1/analytics/org-users for department-level aggregations.")
chart_block("Employee Table",
    "Full list of employees with name, department, job title, employment type, and status.",
    "Central people register for HSE accountability — used to assign CAPA and recognitions.",
    "SELECT * FROM employees WHERE organisation_id=:org LIMIT 200",
    "Table: employees  Columns: id, full_name, department_id, job_title, employment_type, is_active, "
    "organisation_id, email, phone", fill=MID_BLUE)
chart_block("Department Count Cards",
    "KPI cards showing total employees per department.",
    "Helps HR spot understaffed departments and balance HSE resource allocation.",
    "SELECT d.name, COUNT(e.id) FROM employees e JOIN departments d ON e.department_id=d.id "
    "WHERE e.organisation_id=:org GROUP BY d.name",
    "Tables: employees, departments  Columns: department_id, department.name", fill=ACCENT)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5 — RCA
# ══════════════════════════════════════════════════════════════════════════════
section_header("5.  Root Cause Analysis  ( /root-cause-analysis )")
body("Displays all RCA records derived from incidents. API: GET /api/v1/analytics/root-cause-analysis. "
     "Each RCA is built from the incident table — there is no separate RCA table; the backend constructs "
     "RCA objects from incident fields.")
chart_block("RCA List Table",
    "Table of incidents treated as RCA records with ID, incident type, site, zone, conducted-by "
    "(reported_by employee), start/completion dates, root causes, status, priority.",
    "Enables systematic review of why each incident occurred and whether corrective steps were taken.",
    "SELECT i.*, ws.name as station, s.name as site, e.full_name as reporter FROM incidents i "
    "LEFT JOIN working_stations ws ON i.location_station_id=ws.id "
    "LEFT JOIN sites s ON ws.site_id=s.id "
    "LEFT JOIN employees e ON i.reported_by=e.id WHERE i.organisation_id=:org",
    "Tables: incidents, working_stations, sites, employees  Columns: id, incident_type, root_cause, "
    "immediate_cause, investigation_status, report_date, severity",
    notes="Priority is mapped from severity: Critical->Critical, High->High, Medium->Medium, Low->Low. "
    "Status comes from investigation_status field.", fill=DANGER)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6 — ACTIONS (CAPA)
# ══════════════════════════════════════════════════════════════════════════════
section_header("6.  Actions (CAPA)  ( /actions )")
body("Corrective and Preventive Actions. API: GET /api/v1/capa-actions/ (direct model endpoint, "
     "not analytics). Status KPIs computed client-side.")
chart_block("CAPA Table",
    "All CAPA actions with ID, action type, description, responsible person, due date, status.",
    "Tracks every corrective obligation stemming from incidents, near misses, and audits.",
    "SELECT * FROM capa_actions WHERE organisation_id=:org ORDER BY due_date ASC LIMIT 200",
    "Table: capa_actions  Columns: id, action_type, description, responsible_person_id, due_date, "
    "status, organisation_id, incident_id", fill=WARN)
chart_block("Status KPI Cards  (Open / Overdue / Completed)",
    "Three cards showing counts per status band.",
    "Management KPIs for closure rate and overdue backlog.",
    "Open = status NOT IN ('Completed')  Overdue = status != 'Completed' AND due_date < TODAY()  "
    "Completed = status = 'Completed'  — all computed client-side from the full list.",
    "Table: capa_actions  Columns: status, due_date", fill=PURPLE)
chart_block("Overdue Indicator per Row",
    "Each table row shows a badge if due_date < today and status is not Completed.",
    "Draws attention to actions that have passed their deadline.",
    "Client-side: new Date(due_date) < new Date() && status !== 'Completed'",
    "Column: due_date (Date), status (VARCHAR)", fill=DANGER)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7 — COMPLIANCE
# ══════════════════════════════════════════════════════════════════════════════
section_header("7.  Compliance  ( /compliance )")
body("API: GET /api/v1/analytics/compliance-summary. Aggregates data from permits, policies, "
     "incidents, and CAPA to produce the compliance health picture.")
chart_block("Compliance Score  (Big KPI %)",
    "Overall compliance health as a weighted percentage.",
    "Single headline metric for auditors and senior management.",
    "compliance_score = round(\n"
    "  (permit_compliance_pct * 0.4 + policy_review_pct * 0.35 + audit_readiness_pct * 0.25)\n"
    ")\n"
    "where:\n"
    "  permit_compliance_pct = active_permits / max(total_permits,1) * 100\n"
    "  policy_review_pct     = reviewed_policies / max(total_policies,1) * 100\n"
    "  audit_readiness_pct   = 100 - (open_critical_findings / max(total_findings,1) * 100)",
    "Tables: permit_to_works, policys, capa_actions, incidents  Columns: status, review_date, severity",
    fill=TEAL)
chart_block("Legal Register Coverage  (%)",
    "Percentage of policies that have been reviewed at least once.",
    "Regulatory requirement: all legal obligations must be documented and reviewed.",
    "reviewed_policies / total_policies * 100  reviewed = policies WHERE status = 'Active'",
    "Table: policys  Columns: id, status, organisation_id", fill=SUCCESS)
chart_block("Audit Readiness  (%)",
    "Inverse of outstanding critical findings as a fraction of all findings.",
    "High readiness means few unresolved critical issues — the org is ready for an external audit.",
    "audit_readiness_pct = 100 - (open_critical / max(total_findings, 1) * 100)\n"
    "open_critical = capa_actions WHERE status != 'Completed' AND incident.severity IN (1,2)",
    "Tables: capa_actions, incidents  Columns: status, severity", fill=ACCENT)
chart_block("Compliance Trend  (Line Chart — 12 months)",
    "Monthly compliance score for the last 12 months.",
    "Shows whether compliance is improving or eroding over time.",
    "For each month m in last 12: compute compliance_score the same way but scoped to incidents and "
    "permits active in that month.",
    "Tables: incidents, permit_to_works, policys  Columns: report_date, created_at, review_date", fill=MID_BLUE)
chart_block("Findings by Severity  (Donut Chart)",
    "Pie breakdown of CAPA actions by severity of the linked incident.",
    "Visualises whether unresolved work is dominated by high-severity findings.",
    "SELECT i.severity, COUNT(c.id) FROM capa_actions c JOIN incidents i ON c.incident_id=i.id "
    "WHERE c.organisation_id=:org GROUP BY i.severity",
    "Tables: capa_actions, incidents  Columns: incident_id, severity", fill=DANGER)
chart_block("Non-Conformance Table",
    "Top open CAPA actions with action description, owner name, due date, criticality badge.",
    "Provides a concrete action list for compliance officers.",
    "SELECT c.description, e.full_name, c.due_date, i.severity FROM capa_actions c "
    "LEFT JOIN employees e ON c.responsible_person_id=e.id "
    "LEFT JOIN incidents i ON c.incident_id=i.id "
    "WHERE c.status != 'Completed' AND c.organisation_id=:org ORDER BY c.due_date ASC LIMIT 10",
    "Tables: capa_actions, employees, incidents", fill=WARN)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8 — ANALYTICS (7 TABS — contractor tab removed)
# ══════════════════════════════════════════════════════════════════════════════
section_header("8.  Analytics  ( /analytics ) — 7 Tabs")
body("The Analytics page is the most data-rich page. It pulls from three backend endpoints: "
     "/analytics/violations-summary, /analytics/permits-summary, and /analytics/compliance-summary. "
     "Tabs are rendered client-side from the fetched data.")

subhead("Tab 1 — Violations Summary")
chart_block("Incidents by Type  (Bar Chart)",
    "Number of incidents grouped by incident_type (e.g., Near Miss, Injury, Property Damage).",
    "Identifies the most common incident categories to target prevention efforts.",
    "SELECT incident_type, COUNT(*) FROM incidents WHERE org=:org GROUP BY incident_type",
    "Table: incidents  Column: incident_type (VARCHAR)", fill=DANGER)
chart_block("Incidents by Location  (Horizontal Bar)",
    "Incident counts per site/zone.",
    "Identifies geographical hotspots where more safety interventions are needed.",
    "SELECT s.name, COUNT(i.id) FROM incidents i JOIN working_stations ws ON i.location_station_id=ws.id "
    "JOIN sites s ON ws.site_id=s.id WHERE i.organisation_id=:org GROUP BY s.name",
    "Tables: incidents, working_stations, sites  Columns: location_station_id, site_id, name", fill=WARN)
chart_block("Root Cause Donut  (Pie Chart)",
    "Distribution of incidents by root_cause category.",
    "Reveals systemic causes (e.g., human error vs equipment failure) to prioritise training.",
    "SELECT root_cause, COUNT(*) FROM incidents WHERE org=:org GROUP BY root_cause",
    "Table: incidents  Column: root_cause (VARCHAR)", fill=PURPLE)
chart_block("Monthly Incident Trend  (Line — 10 months)",
    "Line chart: incidents per month for the last 10 months.",
    "Core trend indicator; the most-watched chart by safety managers.",
    "For m in last 10 months: SELECT COUNT(*) FROM incidents WHERE MONTH=m AND YEAR=y AND org=:org",
    "Table: incidents  Columns: report_date, organisation_id", fill=ACCENT)
chart_block("Near Miss Monthly Trend  (Dashed Line)",
    "Near-miss count per month, plotted alongside incident trend.",
    "Near misses should outnumber incidents (good reporting culture); if not, it's a warning sign.",
    "For m in last 10 months: SELECT COUNT(*) FROM near_misses WHERE DATE(event_date_time) BETWEEN "
    "start_m AND end_m AND org=:org",
    "Table: near_misses  Columns: event_date_time, organisation_id", fill=MID_BLUE)
chart_block("Injury Category & Person Involved  (Bar Charts)",
    "Two separate bars: one by injury body part/category, one by employment type of person involved.",
    "Reveals whether specific body-part injuries or particular worker groups are disproportionately represented.",
    "Injury cat: SELECT injury_category, COUNT(*) FROM incidents WHERE org=:org GROUP BY injury_category\n"
    "Person: SELECT employment_type, COUNT(*) FROM employees e JOIN incidents i ON i.reported_by=e.id "
    "WHERE org=:org GROUP BY employment_type",
    "Tables: incidents (injury_category), employees (employment_type)", fill=DANGER)
chart_block("Severity Mix  (Stacked Bar — Monthly)",
    "Each bar = one month; stacked by Critical / High / Medium / Low severity.",
    "Shows whether severe incidents are increasing proportionally month-on-month.",
    "For each month m: SELECT severity, COUNT(*) FROM incidents WHERE MONTH=m AND org=:org GROUP BY severity",
    "Table: incidents  Columns: report_date, severity", fill=WARN)

subhead("Tab 2 — Near Miss")
chart_block("Near Miss Table + Trend",
    "List of near-miss records with description, location, date, severity, status.",
    "Near-miss data is a leading indicator — reviewing them prevents future incidents.",
    "SELECT * FROM near_misses WHERE organisation_id=:org ORDER BY event_date_time DESC LIMIT 200",
    "Table: near_misses  Columns: id, description, location_station_id, event_date_time, "
    "potential_consequence, underlying_cause, reported_by, capa_escalation", fill=TEAL)

subhead("Tab 3 — RCA (Root Cause)")
chart_block("Root Cause Analysis List",
    "Same RCA data as /root-cause-analysis page but shown as a tab inside Analytics.",
    "Allows analysts to cross-reference RCA findings alongside other analytics.",
    "GET /api/v1/analytics/root-cause-analysis  (same endpoint as RCA page)",
    "Tables: incidents, working_stations, sites, employees", fill=PURPLE)

subhead("Tab 4 — Permits & Active Work")
chart_block("Active Permits KPI + Radar Chart",
    "Total active permits, total workers on site, risk-assessment radar by work type.",
    "Tracks live work permits to ensure all active work is authorised.",
    "active_permits = SELECT COUNT(*) FROM permit_to_works WHERE status='Active' AND org=:org\n"
    "workers = SELECT SUM(number_of_workers) FROM permit_to_works WHERE status='Active' AND org=:org\n"
    "radar = SELECT pt.name, COUNT(p.id) FROM permit_to_works p JOIN permit_types pt "
    "ON p.permit_type_id=pt.id WHERE org=:org GROUP BY pt.name",
    "Tables: permit_to_works, permit_types  Columns: status, number_of_workers, permit_type_id, name",
    fill=ACCENT)
chart_block("Permit Violations Feed",
    "Recent violations: permits that expired or were used without sign-off.",
    "Provides an audit trail of permit non-compliance events.",
    "SELECT p.id, p.expiry_date FROM permit_to_works p WHERE p.expiry_date < TODAY() AND p.status='Active' "
    "AND org=:org ORDER BY expiry_date DESC LIMIT 5",
    "Table: permit_to_works  Columns: expiry_date, status", fill=DANGER)
chart_block("Expiry Timeline  (Gantt-style Bars)",
    "Horizontal bars showing how many days remain before each active permit expires.",
    "Allows permit coordinators to prioritise renewals.",
    "For each active permit: days_left = expiry_date - TODAY()\n"
    "bar_width% = days_left / max_days * 100",
    "Table: permit_to_works  Columns: id, expiry_date, permit_type_id", fill=WARN)
chart_block("Permit Status by Type  (Stacked Bar)",
    "For each permit type: stacked bar showing count of Active / Expired / Closed permits.",
    "Reveals which work types have high expiry rates — a risk signal.",
    "SELECT pt.name, p.status, COUNT(*) FROM permit_to_works p JOIN permit_types pt "
    "ON p.permit_type_id=pt.id WHERE p.organisation_id=:org GROUP BY pt.name, p.status",
    "Tables: permit_to_works, permit_types  Columns: permit_type_id, status, name", fill=ACCENT)

subhead("Tab 5 — Zone Risk")
chart_block("Incident Count by Zone  (Bar Chart)",
    "Incidents grouped by working_station zone for the org.",
    "Pinpoints physical areas with the highest risk concentration.",
    "SELECT ws.zone_name, COUNT(i.id) FROM incidents i JOIN working_stations ws "
    "ON i.location_station_id=ws.id WHERE i.organisation_id=:org GROUP BY ws.zone_name",
    "Tables: incidents, working_stations  Columns: location_station_id, zone_name", fill=DANGER)
chart_block("Zone Risk Summary Table",
    "Table: zone name, incident count, severity badge (High/Medium/Low).",
    "Quick reference for which zones need immediate physical safety improvements.",
    "severity badge: count > 10 -> High (red), 5-10 -> Medium (amber), < 5 -> Low (green)",
    "Tables: incidents, working_stations  Columns: zone_name, severity", fill=WARN)

subhead("Tab 6 — Trend Reports")
chart_block("Monthly Incidents vs Near Misses  (Dual Line Chart)",
    "Two lines on the same axis: monthly incidents and monthly near-misses for last 10 months.",
    "The ratio of near-misses to incidents indicates reporting culture health "
    "(Heinrich Triangle principle: more near-miss reports = better culture).",
    "Incidents line: same as violations tab monthly trend.\n"
    "Near-miss line: same as near-miss monthly trend.",
    "Tables: incidents (report_date), near_misses (event_date_time)", fill=ACCENT)
chart_block("Severity Mix  (Stacked Bar — Monthly)",
    "Same stacked-bar chart as Violations tab but placed here for trend context.",
    "Allows managers to see severity escalation month-on-month at a glance.",
    "Same computation as Violations tab severity mix.",
    "Table: incidents  Columns: report_date, severity", fill=MID_BLUE)

subhead("Tab 7 — PPE Compliance")
note("Status: Empty state shown — 'No PPE compliance data available'. No dedicated PPE table exists in "
     "the database. This tab is reserved for future data collection.")
subhead("Customer Reports (Scheduled Reports)")
note("Status: Empty state shown — 'No scheduled reports configured'. No scheduled-report table exists "
     "in the database. Reserved for future reporting automation.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9 — RISK
# ══════════════════════════════════════════════════════════════════════════════
section_header("9.  Risk  ( /risk )")
body("Three API calls: /analytics/risk-summary, /analytics/residual-risk-trend, /analytics/risk-matrix. "
     "Covers risk matrix heatmap, zone risk bar chart, active task table, and risk aging.")
chart_block("Control Effectiveness  (KPI %)",
    "Percentage of hazards that have at least one linked CAPA action (indicating a control is in place).",
    "Measures how thoroughly identified hazards are being mitigated.",
    "controlled_hazards = COUNT(DISTINCT hazard_id FROM capa_actions WHERE org=:org)\n"
    "total_hazards = COUNT(*) FROM hazards WHERE org=:org\n"
    "control_effectiveness = controlled_hazards / max(total_hazards, 1) * 100  -> e.g. '73%'",
    "Tables: hazards, capa_actions  Columns: id, organisation_id, hazard_id", fill=SUCCESS)
chart_block("Unverified Controls  (KPI Count)",
    "Number of CAPA actions with status = 'Pending' — controls assigned but not yet verified.",
    "Unverified controls represent gaps where a fix was planned but not confirmed effective.",
    "SELECT COUNT(*) FROM capa_actions WHERE status='Pending' AND organisation_id=:org",
    "Table: capa_actions  Column: status", fill=WARN)
chart_block("Risk Escalations  (KPI Count)",
    "Count of incidents with severity = Critical (value 1) reported in the last 90 days.",
    "Critical incidents that are recent represent elevated organisational risk.",
    "SELECT COUNT(*) FROM incidents WHERE severity=1 AND organisation_id=:org AND report_date >= TODAY() - 90 days",
    "Table: incidents  Columns: severity, report_date, organisation_id", fill=DANGER)
chart_block("Zone Risk  (Horizontal Bar Chart)",
    "Incident count per working_station zone, displayed as proportional horizontal bars.",
    "Visualises which physical zones accumulate the most risk.",
    "SELECT ws.zone_name AS zone, COUNT(i.id) AS value FROM incidents i "
    "JOIN working_stations ws ON i.location_station_id=ws.id "
    "WHERE i.organisation_id=:org GROUP BY ws.zone_name ORDER BY value DESC",
    "Tables: incidents, working_stations  Columns: location_station_id, zone_name",
    notes="Progress bar width is computed as: zone_value / max_zone_value * 100 — showing relative dominance.",
    fill=TEAL)
chart_block("Residual Risk Trend  (Area / Line Chart — Quarterly)",
    "Line chart showing average incident severity per quarter for the last 4 quarters.",
    "Residual risk is what remains after controls are applied; a downward trend = controls working.",
    "For each quarter q: SELECT AVG(CAST(severity AS FLOAT)) FROM incidents "
    "WHERE report_date BETWEEN q_start AND q_end AND organisation_id=:org\n"
    "Quarter labels: Q1...Q4 of current year",
    "Table: incidents  Columns: report_date, severity (1=Critical...4=Low)", fill=ACCENT)
chart_block("Risk Matrix  (5x5 Heatmap Grid)",
    "5-column x 5-row coloured grid. Each cell shows count of incidents at that Likelihood x Severity intersection.",
    "Standard HSE risk assessment tool; the top-right quadrant (high likelihood + high severity) is the danger zone.",
    "grid[row][col] = COUNT(*) FROM incidents WHERE severity_bucket=row AND likelihood_bucket=col AND organisation_id=:org\n"
    "Severity buckets: severity in {1,2,3,4,5} maps directly to row index.\n"
    "Likelihood is approximated from injury_category / incident recurrence.",
    "Table: incidents  Columns: severity, injury_category, report_date", fill=DANGER)
chart_block("Active Risk Tasks Table",
    "Table of open CAPA actions linked to incidents, with description, owner, due date, status.",
    "Provides a concrete task list for risk owners to action.",
    "SELECT c.id, c.description, e.full_name AS owner, c.due_date, c.status FROM capa_actions c "
    "LEFT JOIN employees e ON c.responsible_person_id=e.id "
    "WHERE c.status != 'Completed' AND c.organisation_id=:org ORDER BY c.due_date ASC LIMIT 10",
    "Tables: capa_actions, employees  Columns: description, responsible_person_id, due_date, status", fill=MID_BLUE)
chart_block("Risk Aging  (Grouped Bar Chart)",
    "Bars grouped by age buckets (0-30 days, 31-60, 61-90, 91+ days). Each bar stacked by severity "
    "(Low / Medium / High / Critical).",
    "Older open risks are more dangerous; this chart pressure-tests the closure velocity.",
    "For each incident i: age_days = TODAY() - i.report_date\n"
    "bucket: 0-30, 31-60, 61-90, 91+\n"
    "severity: from incidents.severity\n"
    "-> aggregated as {bucket, low, medium, high, critical, line}",
    "Table: incidents  Columns: report_date, severity, organisation_id", fill=WARN)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10 — EQUIPMENT CERTIFICATION
# ══════════════════════════════════════════════════════════════════════════════
section_header("10.  Equipment Certification  ( /equipment-certification )")
body("API: GET /api/v1/equipment-certifications/. Dedicated table with its own controller. "
     "Tracks certification status of all safety-critical equipment.")
chart_block("Certification Table",
    "Full list of equipment records: name, type, serial number, manufacturer, model, certification type, "
    "issue date, expiry date, next inspection date, status, compliance standard.",
    "Regulatory compliance requires proof that all safety equipment is certified and within validity.",
    "SELECT * FROM equipment_certifications WHERE organisation_id=:org ORDER BY expiry_date ASC",
    "Table: equipment_certifications  Columns: equipment_name, equipment_type, serial_number, manufacturer, "
    "model, certification_type, certified_by, issue_date, expiry_date, next_inspection_date, status, "
    "compliance_standard, organisation_id", fill=MID_BLUE)
chart_block("Status KPI Cards  (Valid / Expiring / Expired)",
    "Three KPI cards: count of certs in each validity band.",
    "Allows maintenance teams to see at a glance how many certs need urgent renewal.",
    "Valid    = status = 'Active' AND expiry_date > TODAY()\n"
    "Expiring = status = 'Active' AND expiry_date BETWEEN TODAY() AND TODAY()+30\n"
    "Expired  = status = 'Expired' OR expiry_date < TODAY()",
    "Table: equipment_certifications  Columns: status, expiry_date", fill=TEAL)
chart_block("Equipment Type Breakdown  (Donut)",
    "Pie/donut showing how many certifications exist per equipment_type.",
    "Identifies which equipment categories are most heavily regulated.",
    "SELECT equipment_type, COUNT(*) FROM equipment_certifications WHERE organisation_id=:org GROUP BY equipment_type",
    "Table: equipment_certifications  Column: equipment_type", fill=ACCENT)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11 — VIOLATIONS
# ══════════════════════════════════════════════════════════════════════════════
section_header("11.  Violations  ( /violations )  &  Detail Page")
body("Two sub-pages: the list page and the detail page. List: GET /api/v1/incidents/?limit=10. "
     "Detail: GET /api/v1/analytics/violation-detail/{incident_id}.")
chart_block("Violations Summary  (Analytics Charts — top of page)",
    "Charts from /analytics/violations-summary: incidents by type, by location, root cause donut, monthly trend.",
    "See Section 8 Tab 1 for full formula details.",
    "GET /api/v1/analytics/violations-summary",
    "Tables: incidents, working_stations, sites", fill=DANGER)
chart_block("Recent Incidents Table  (bottom of list page)",
    "Table of the 10 most recent incidents with ID, type, severity badge, status, date, zone. "
    "Rows are clickable and navigate to the detail page.",
    "Provides direct access to drill into specific incidents.",
    "SELECT id, incident_type, severity, investigation_status, report_date, location_station_id "
    "FROM incidents WHERE organisation_id=:org ORDER BY report_date DESC LIMIT 10",
    "Table: incidents  Columns: id, incident_type, severity, investigation_status, report_date", fill=WARN)

subhead("Violation Detail Page  ( /violations/:id )")
chart_block("Incident Header  (ID, Severity, Status)",
    "Displays formatted incident ID (e.g., INC-00042), severity badge, investigation status.",
    "At-a-glance identification of the specific incident.",
    "id formatted as: INC-{id:05d}\n"
    "severity mapped: 1->Critical (red), 2->High (orange), 3->Medium (yellow), 4->Low (green)\n"
    "status: investigation_status field directly from DB",
    "Table: incidents  Columns: id, severity, investigation_status", fill=DANGER)
chart_block("Status Tracker  (Step Progress Bar 0-4)",
    "5-step horizontal stepper: Reported -> Under Review -> Investigation -> CAPA -> Closed.",
    "Shows where in the investigation lifecycle this incident currently sits.",
    "status_step = {Pending:0, 'Under Investigation':1, 'In Progress':2, Completed:3, Closed:4}.get(investigation_status, 0)",
    "Table: incidents  Column: investigation_status", fill=ACCENT)
chart_block("Details Grid  (9 fields)",
    "Zone, Site, Station, Reporter name, Timestamp, Persons Involved, Days Away, Permit Active, Control Failure.",
    "Full factual record of the incident for investigators.",
    "SELECT i.*, ws.name AS station, s.name AS site, e.full_name AS reporter FROM incidents i "
    "LEFT JOIN working_stations ws ON i.location_station_id=ws.id "
    "LEFT JOIN sites s ON ws.site_id=s.id "
    "LEFT JOIN employees e ON i.reported_by=e.id WHERE i.id=:incident_id AND i.organisation_id=:org",
    "Tables: incidents, working_stations, sites, employees", fill=MID_BLUE)
chart_block("CAPA Actions Table  (on detail page)",
    "All CAPA actions linked to this incident: action type, description, responsible person, due date, status.",
    "Every incident should have corrective actions attached; this table shows the full CAPA trail.",
    "SELECT c.*, e.full_name AS responsible_person FROM capa_actions c "
    "LEFT JOIN employees e ON c.responsible_person_id=e.id "
    "WHERE c.incident_id=:incident_id AND c.organisation_id=:org",
    "Tables: capa_actions, employees", fill=PURPLE)
chart_block("Event Timeline",
    "Chronological log of actions: Incident Reported, Investigation Started, CAPA Created, CAPA Completed, "
    "Closed — with timestamps.",
    "Provides an audit trail for HSE investigators and regulators.",
    "Timeline is constructed from: incident.report_date (Reported), first capa_action.created_at (CAPA Created), "
    "capa_action where status=Completed (Completed), incident.investigation_status=Closed (Closed)",
    "Tables: incidents, capa_actions  Columns: report_date, created_at, status", fill=TEAL)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12 — ENGAGEMENT
# ══════════════════════════════════════════════════════════════════════════════
section_header("12.  Engagement  ( /engagement )")
body("API: GET /api/v1/analytics/engagement-summary. Aggregates safety participation metrics from "
     "safety_walks, incidents, near_misses, employees, capa_actions, and sites.")
chart_block("Reporting Rate  (Big KPI %)",
    "Percentage of employees who submitted at least one incident or near-miss report this month.",
    "Measures safety reporting culture — low rate = under-reporting (dangerous).",
    "this_month_reports = incidents_this_month + near_misses_this_month\n"
    "reporting_rate = min(100, round(this_month_reports / total_employees * 25))\n\n"
    "Note: x25 scaling because each report by one person represents roughly 25% participation proxy "
    "(tunable constant).\n"
    "Fallback if no employees: min(100, reports * 10)",
    "Tables: incidents (report_date), near_misses (event_date_time), employees (id)", fill=TEAL)
chart_block("Reporting Rate MoM Arrow  (up / down)",
    "Arrow indicating whether this month's report count is higher or lower than last month.",
    "Trend direction tells managers if engagement is improving.",
    "reporting_rate_mom = this_month_reports - last_month_reports\n"
    "up if >= 0,  down if < 0",
    "Tables: incidents, near_misses  (counts scoped to this_month vs last_month date windows)", fill=SUCCESS)
chart_block("Engagement Survey Score  (N.N / 5)",
    "Average compliance_rating across all safety walks for the org.",
    "Proxy for how well safety procedures are being followed; maps the '5-star' model to a meaningful "
    "engagement score.",
    "survey_score = AVG(compliance_rating) FROM safety_walks WHERE org=:org\n"
    "survey_score_pct = round(survey_score / 5 * 100)",
    "Table: safety_walks  Column: compliance_rating (Integer 1-5)", fill=ACCENT)
chart_block("Survey Score MoM Arrow  (up / down — hidden if no prior data)",
    "Direction indicator comparing this month's avg compliance_rating to last month's.",
    "Reveals whether safety walk quality is improving.",
    "avg_this_month = AVG(compliance_rating) WHERE inspection_date_time >= first_of_this_month\n"
    "avg_last_month = AVG(compliance_rating) WHERE date BETWEEN first_of_last_month AND first_of_this_month\n"
    "survey_score_mom = round(avg_this_month - avg_last_month, 1)\n"
    "Arrow shown only when BOTH months have data; hidden (null) otherwise.",
    "Table: safety_walks  Columns: compliance_rating, inspection_date_time", fill=MID_BLUE)
chart_block("Safety Observations Ring  (%)",
    "Donut ring: % of safety walks with high compliance (rating >= 4).",
    "High-rated walks indicate inspectors found good safety conditions — a positive metric.",
    "compliant_walks = COUNT(*) FROM safety_walks WHERE compliance_rating >= 4 AND org=:org\n"
    "total_walks = COUNT(*) FROM safety_walks WHERE org=:org\n"
    "safety_observations_pct = round(compliant_walks / max(total_walks, 1) * 100)",
    "Table: safety_walks  Column: compliance_rating", fill=SUCCESS)
chart_block("Safety Walks Ring  (%)",
    "Donut ring: % of employees who completed at least one safety walk this month.",
    "Participation in safety walks is a leading HSE activity indicator.",
    "walks_this_month = COUNT(*) FROM safety_walks WHERE inspection_date_time >= first_of_month AND org=:org\n"
    "effective_walks = walks_this_month if > 0 else total_walks  (fallback for historical data)\n"
    "safety_walks_pct = min(100, round(effective_walks / max(total_employees, 1) * 100))",
    "Tables: safety_walks (inspection_date_time), employees (id)", fill=TEAL)
chart_block("Toolbox Attendance Ring  (%)",
    "Donut ring: % of employees who attended a toolbox-type safety walk.",
    "Toolbox talks are mandatory pre-work briefings; attendance tracks compliance.",
    "toolbox_count = COUNT(*) FROM safety_walks WHERE LOWER(inspection_type) LIKE '%toolbox%' AND org=:org\n"
    "fallback if toolbox_count=0: use total_walks (no toolbox entries in DB)\n"
    "toolbox_pct = min(100, round(toolbox_count / max(total_employees, 1) * 100))",
    "Table: safety_walks  Columns: inspection_type, organisation_id", fill=ACCENT)
chart_block("Site Participation Ring  (%)",
    "Donut ring: % of sites that had at least one incident or near-miss reported.",
    "Sites with no reports may have under-reporting problems or genuine zero-incident performance.",
    "active_sites = MAX(DISTINCT site_ids from incidents this org, DISTINCT site_ids from near_misses this org)\n"
    "total_sites = COUNT(*) FROM sites WHERE org=:org\n"
    "site_participation_pct = round(active_sites / max(total_sites, 1) * 100)",
    "Tables: incidents, near_misses, working_stations, sites", fill=PURPLE)
chart_block("Top Recognitions  (3 Employee Avatars)",
    "The 3 employees with the most completed CAPA actions, displayed as recognition cards.",
    "Positive reinforcement for safety champions — promotes a safety-first culture.",
    "SELECT e.full_name, COUNT(c.id) AS cnt FROM employees e JOIN capa_actions c "
    "ON c.responsible_person_id=e.id WHERE c.status='Completed' AND e.organisation_id=:org "
    "GROUP BY e.full_name ORDER BY cnt DESC LIMIT 3\n\n"
    "Fallback: if no completed CAPAs, show top 3 non-admin org users by ID.",
    "Tables: employees, capa_actions  Columns: full_name, responsible_person_id, status", fill=WARN)
chart_block("Open Actions List  (Checklist with Status Pills)",
    "Up to 5 oldest open CAPA actions with status pills: Due Today / Due Tomorrow / Overdue.",
    "Brings the most urgent safety obligations directly onto the engagement page.",
    "SELECT c.description, c.action_type, c.due_date FROM capa_actions c "
    "WHERE c.status != 'Completed' AND c.organisation_id=:org "
    "ORDER BY CASE WHEN due_date IS NULL THEN 1 ELSE 0 END, due_date ASC LIMIT 5\n\n"
    "Status pill logic:\n"
    "  days = due_date - TODAY()\n"
    "  days < 0  -> 'Overdue'\n"
    "  days == 0 -> 'Due Today'\n"
    "  else      -> 'Due Tomorrow'",
    "Table: capa_actions  Columns: description, action_type, due_date, status", fill=DANGER)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13 — FORMULA CHEAT SHEET
# ══════════════════════════════════════════════════════════════════════════════
section_header("13.  Formula Cheat-Sheet")
body("Quick reference for every computed metric in the platform.")
formulas = [
    ["Metric", "Formula", "Output"],
    ["Compliance Score", "(permit_compliance x0.4 + policy_review x0.35 + audit_readiness x0.25)", "%  0-100"],
    ["Permit Compliance %", "active_permits / total_permits x 100", "%  0-100"],
    ["Policy Review %", "active_policies / total_policies x 100", "%  0-100"],
    ["Audit Readiness %", "100 - (open_critical_findings / total_findings x 100)", "%  0-100"],
    ["Control Effectiveness %", "controlled_hazards / total_hazards x 100", "%  0-100"],
    ["Reporting Rate %", "min(100, (incidents+nearmisses this month) / employees x 25)", "%  0-100"],
    ["Reporting Rate MoM", "this_month_count - last_month_count", "Integer (pos/neg)"],
    ["Survey Score", "AVG(safety_walks.compliance_rating)", "Float 0-5"],
    ["Survey Score %", "survey_score / 5 x 100", "%  0-100"],
    ["Survey Score MoM", "avg_this_month_rating - avg_last_month_rating", "Float (pos/neg/null)"],
    ["Safety Observations %", "walks_rating>=4 / total_walks x 100", "%  0-100"],
    ["Safety Walks %", "min(100, effective_walks / employees x 100)", "%  0-100"],
    ["Toolbox Attendance %", "min(100, toolbox_walks / employees x 100)", "%  0-100"],
    ["Site Participation %", "max(active_sites_inc, active_sites_nm) / total_sites x 100", "%  0-100"],
    ["Zone Risk Bar Width", "zone_value / max_zone_value x 100", "%  0-100 (CSS width)"],
    ["Risk Matrix Cell", "COUNT(incidents) at severity_bucket x likelihood_bucket", "Integer"],
    ["Residual Risk Trend pt", "AVG(severity) per quarter  (1=Critical...4=Low)", "Float 1-4"],
    ["Risk Aging Bucket", "TODAY() - incidents.report_date -> 0-30/31-60/61-90/91+", "Days bucket"],
    ["Equipment Status", "expiry_date vs TODAY(): Valid / Expiring(<30d) / Expired", "Label"],
    ["CAPA Status Pill", "due_date-TODAY(): <0->Overdue, =0->Due Today, >0->Due Tomorrow", "Label"],
    ["Severity Mapping", "1->Critical, 2->High, 3->Medium, 4->Low", "Label + colour"],
]
grid_table(formulas, [50, FULL_W - 82, 32], first_col_bold=True)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14 — DB TABLE REFERENCE  (contractor references removed)
# ══════════════════════════════════════════════════════════════════════════════
section_header("14.  DB Table Reference")
body("Every table in the MySQL database, its key columns, and which pages/charts consume it.")
db_tables = [
    ["Table Name", "Key Columns", "Used By"],
    ["incidents", "id, organisation_id, report_date, incident_type, severity (1-4), investigation_status, "
     "root_cause, immediate_cause, description, location_station_id, reported_by, injury_category, "
     "number_persons_involved, days_away_from_work, permit_active, control_failure",
     "Dashboard, Analytics (all tabs), Risk, Violations, RCA, Compliance, Engagement"],
    ["near_misses", "id, organisation_id, event_date_time, description, potential_consequence, "
     "underlying_cause, reported_by, location_station_id, capa_escalation",
     "Dashboard, Analytics (Near Miss tab), Engagement (reporting rate)"],
    ["capa_actions", "id, organisation_id, incident_id, action_type, description, responsible_person_id, "
     "due_date, status",
     "Actions page, Compliance, Risk (active tasks), Engagement (open actions, recognitions), Violation detail"],
    ["safety_walks", "id, organisation_id, inspection_date_time, inspector_id, inspection_type, "
     "compliance_rating (1-5), issues_found, critical_issues, follow_up_required, location_station_id",
     "Checklists, Engagement (all ring metrics + survey score)"],
    ["employees", "id, organisation_id, full_name, department_id, job_title, employment_type, is_active, "
     "email, phone",
     "Users page, Engagement (reporting rate, recognitions), CAPA owner names, Violation detail (reporter)"],
    ["permit_to_works", "id, organisation_id, permit_type_id, status, expiry_date, number_of_workers",
     "Analytics Permits tab, Compliance (permit compliance %)"],
    ["permit_types", "id, name", "Analytics Permits tab (permit type labels)"],
    ["hazards", "id, organisation_id, description, category_id", "Risk page (Control Effectiveness KPI)"],
    ["equipment_certifications", "id, organisation_id, equipment_name, equipment_type, serial_number, "
     "manufacturer, model, certification_type, certified_by, issue_date, expiry_date, next_inspection_date, "
     "status, compliance_standard",
     "Equipment Certification page (table, KPI cards, type donut)"],
    ["policys", "id, organisation_id, policy_name, category, issue_date, owner, status",
     "Compliance (policy review %, legal register coverage)"],
    ["sites", "id, organisation_id, name",
     "Dashboard (heatmap), Analytics (by location), Engagement (site participation)"],
    ["working_stations", "id, site_id, name, zone_name",
     "Dashboard, Analytics, Risk (zone charts), Violations (station name), RCA"],
    ["departments", "id, organisation_id, name", "Users page (department breakdown cards)"],
    ["organisations", "id, name, subscription_plan",
     "Multi-tenant root — every other table references this via organisation_id"],
    ["users", "id, organisation_id, username, email, full_name, password_hash, app_role_id, is_active",
     "Auth (login, JWT), Engagement (fallback recognitions if employees empty)"],
]
grid_table(db_tables, [40, 74, FULL_W - 114], first_col_bold=True)
spacer(4)
end = doc.add_paragraph(); end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = end.add_run("End of Document.  Generated from the live source code of hse_old_ui. "
                "All formulas reflect the actual Python/SQL logic in backend/app/controllers/analytics.py. "
                "All DB table references verified against backend/app/models/.")
r.font.size = Pt(8); r.font.color.rgb = RGBColor.from_string(TEXT_MID)

doc.save(OUTPUT)
print(f"DOCX saved to: {OUTPUT}")
