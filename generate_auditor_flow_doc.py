"""
WF-05 Auditor Flow — narrative Word document generator.

Tells the story of one real audit end to end. The numbers are not illustrative:
AUD-000010 is an audit that was actually run through this system against the live
database, and every score, classification and timestamp quoted here is what the
software produced.
"""
import os
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette ───────────────────────────────────────────────────────────────────
ORANGE, DARK, MID = "EA580C", "0F172A", "334155"
BLUE, AMBER, RED, GREEN, PURPLE = "2563EB", "B45309", "DC2626", "047857", "7C3AED"
LIGHT_BG, GREY, TEXT_DARK, TEXT_MID = "FFF7ED", "F1F5F9", "0F172A", "475569"
WHITE = "FFFFFF"
FULL_W = 174

doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Mm(16)
    s.left_margin = s.right_margin = Mm(18)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9.5)
style.paragraph_format.space_after = Pt(4)


# ── Helpers ───────────────────────────────────────────────────────────────────

def shade(cell, hexfill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexfill)
    cell._tc.get_or_add_tcPr().append(el)


def fill_cell(cell, text, bold=False, color=None, size=9, align=None, mono=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = "Consolas" if mono else "Calibri"
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    if fill:
        shade(cell, fill)


def set_widths(table, widths_mm):
    """Pin column widths.

    Setting cell.width alone is not enough: Word's autofit algorithm recomputes
    the layout from the content and ignores it, which is why a 24mm label column
    renders half the page wide. The fixed tblLayout element is what makes the
    widths authoritative.
    """
    table.autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)

    # Under a fixed layout Word reads the tblGrid, not the cell widths, so both
    # have to be set. Setting only the cells is why a 24mm label column still
    # rendered nearly half the page wide.
    grid_el = table._tbl.find(qn("w:tblGrid"))
    if grid_el is not None:
        for col, w in zip(grid_el.findall(qn("w:gridCol")), widths_mm):
            col.set(qn("w:w"), str(int(Mm(w).twips)))

    for row in table.rows:
        for i, w in enumerate(widths_mm):
            if i < len(row.cells):
                row.cells[i].width = Mm(w)


def spacer(pts=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.space_before = Pt(0)
    return p


def part_header(number, title, subtitle=None):
    """A major chapter break."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t.rows[0].cells[0]
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{number}   {title}")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(WHITE)
    shade(c, ORANGE)
    set_widths(t, [FULL_W])
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(4)
        r2 = p2.add_run(subtitle)
        r2.italic = True
        r2.font.size = Pt(9.5)
        r2.font.color.rgb = RGBColor.from_string(TEXT_MID)
    spacer(3)


def subhead(text, color=DARK, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(9)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def body(text, bold=False, size=9.5, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color or TEXT_DARK)
    return p


def bullet(text, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    if bold_lead:
        r = p.add_run(bold_lead)
        r.bold = True
        r.font.size = Pt(9.5)
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    return p


def quote(text, source=None):
    """A line lifted from the specification."""
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f"“{text}”")
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(MID)
    if source:
        r2 = p.add_run(f"   — {source}")
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor.from_string("94A3B8")
    shade(c, GREY)
    set_widths(t, [FULL_W])
    spacer(4)


def callout(title, text, colour=BLUE, fill="EFF6FF"):
    t = doc.add_table(rows=1, cols=1)
    c = t.rows[0].cells[0]
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(colour)
    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    r2 = p2.add_run(text)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor.from_string(TEXT_MID)
    shade(c, fill)
    set_widths(t, [FULL_W])
    spacer(5)


def screen_note(where, text):
    """What the person actually sees, and on which surface."""
    t = doc.add_table(rows=1, cols=2)
    fill_cell(t.rows[0].cells[0], where.upper(), bold=True, color=WHITE, size=8,
              align=WD_ALIGN_PARAGRAPH.CENTER,
              fill=ORANGE if where.lower().startswith("mobile") else BLUE)
    fill_cell(t.rows[0].cells[1], text, size=9, color=TEXT_MID, fill="FFFFFF")
    set_widths(t, [24, FULL_W - 24])
    spacer(5)


def grid(data, widths, header_fill=DARK, first_col_bold=False, small=8.5):
    t = doc.add_table(rows=len(data), cols=len(data[0]))
    t.style = "Table Grid"
    for i, row in enumerate(data):
        for j, val in enumerate(row):
            if i == 0:
                fill_cell(t.rows[i].cells[j], val, bold=True, color=WHITE, size=small, fill=header_fill)
            else:
                fill_cell(t.rows[i].cells[j], val, bold=(first_col_bold and j == 0),
                          size=small, fill=WHITE if i % 2 else "FAFBFC")
    set_widths(t, widths)
    spacer(6)


# ══════════════════════════════════════════════════════════════════════════════
# Cover
# ══════════════════════════════════════════════════════════════════════════════

t = doc.add_table(rows=1, cols=1)
c = t.rows[0].cells[0]
c.text = ""
p = c.paragraphs[0]
p.paragraph_format.space_before = Pt(10)
r = p.add_run("WF-05 · AUDIT, INSPECTION & COMPLIANCE MONITORING")
r.bold = True
r.font.size = Pt(9)
r.font.color.rgb = RGBColor.from_string("FED7AA")

p2 = c.add_paragraph()
r2 = p2.add_run("The Complete Auditor Flow")
r2.bold = True
r2.font.size = Pt(26)
r2.font.color.rgb = RGBColor.from_string(WHITE)

p3 = c.add_paragraph()
r3 = p3.add_run("Schedule to Verified Closure")
r3.font.size = Pt(15)
r3.font.color.rgb = RGBColor.from_string("FFEDD5")

p4 = c.add_paragraph()
p4.paragraph_format.space_after = Pt(12)
r4 = p4.add_run(
    "The story of one audit, from the calendar that booked it to the day its last "
    "corrective action was verified on site."
)
r4.italic = True
r4.font.size = Pt(10)
r4.font.color.rgb = RGBColor.from_string("FFEDD5")
shade(c, ORANGE)
set_widths(t, [FULL_W])
spacer(10)

grid([
    ["Sources", "EHSERA-ISMS-AO-2026-v1.0 · ALGO-ISMS-WF-2026-v1.0 · AUD-FORM-01"],
    ["Standards", "ISO 45001 Cl.9 · ISO 14001 Cl.9 · OSHA VPP"],
    ["Surfaces", "Mobile (React Native) — the audit. Web console (React) — before and after."],
    ["Roles", "Auditor · Worker · Supervisor · Safety Manager (mobile + web) · Admin (web only) · The System"],
    ["Worked example", "AUD-000010 — Q3 Fire Safety Audit, WindTech Nacelle Manufacturing Ltd"],
], [30, FULL_W - 30], header_fill=DARK, first_col_bold=True, small=9)

callout(
    "Every number in this document is real.",
    "AUD-000010 was run end to end through the built system against the live database. The score, "
    "the classifications, the escalations that fired and the timestamps are what the software "
    "actually produced — not an illustration of what it might produce.",
    colour=GREEN, fill="ECFDF5",
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# The premise
# ══════════════════════════════════════════════════════════════════════════════

part_header("", "Before the story starts", "What kind of thing an audit is, in this system.")

body(
    "An audit here is not a form somebody fills in. It is a ten-step process with named owners, "
    "hard stops that cannot be skipped, and a set of rules that fire on their own without anyone "
    "having to notice. Four things follow from that, and they shape everything below."
)

subhead("Nobody books an audit by hand")
quote(
    "Audits are not booked by hand. The system generates the annual programme from each site's risk "
    "band, and that band is driven by the site's own safety performance score. A site that "
    "deteriorates gets audited more often, automatically."
)
body(
    "The calendar is generated from data the platform already holds. A site whose safety performance "
    "score climbs into the critical band starts receiving monthly inspections and quarterly audits "
    "without a person deciding it should."
)

subhead("The audit happens on the phone, in the field, as it happens")
quote("Every observation logged live on the app — no paper, no writing up afterwards.")
body(
    "Steps 4 to 8 run on the phone while the auditor is standing in the place they are auditing. "
    "There is no write-up stage afterwards, because a write-up stage is where detail goes to die and "
    "where the gap between what was seen and what was recorded opens up."
)

subhead("The auditor judges; the system does the arithmetic")
body(
    "The auditor decides what each finding is. The system calculates the score, applies the "
    "thresholds, flags repeats, and raises the corrective actions. Neither does the other's job — "
    "and two rules are taken out of the auditor's hands entirely, which we come to at step 7."
)

subhead("Issuing the report is not the end")
quote(
    "An audit is not closed when the report is issued — it stays open until every corrective action "
    "it raised has been verified effective."
)
body(
    "This is the sentence that most audit software gets wrong. Marking an action complete is a claim "
    "made by whoever owned it. Verification is somebody standing in the same place afterwards, "
    "checking the fix is holding. The two are different events and this system keeps them apart."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# Step 01
# ══════════════════════════════════════════════════════════════════════════════

part_header("STEP 01", "The calendar builds itself",
            "PLAN · The System · automatic — no human involved")

body(
    "The story begins with nobody. On the audit programme screen, the platform has already scored "
    "every site in the organisation and sorted them into four risk bands. WindTech Nacelle "
    "Manufacturing Ltd sits at a safety performance score of 43.1, which places it in the medium "
    "band: quarterly inspection, bi-annual audit."
)

body("The band is not a label somebody typed. It comes from the same KPI engine the safety "
     "dashboards read, so a site cannot be critical on one screen and low on another.")

grid([
    ["Band", "Qualifying criteria", "How often", "Re-audit trigger"],
    ["CRITICAL", "Score 75+ or any fatal/critical event in 12 months",
     "Monthly inspection + quarterly full audit", "Any Major NC → re-audit within 30 days"],
    ["HIGH", "Score 50–74 or a lost-time injury in 6 months",
     "Monthly inspection + quarterly audit", "Two or more Major NCs in 12 months → 60 days"],
    ["MEDIUM", "Score 25–49", "Quarterly inspection + bi-annual audit",
     "Overall score below 65% for two consecutive audits"],
    ["LOW", "Score 0–24, no serious event in 24 months", "Bi-annual inspection + annual audit",
     "Any lost-time injury or score below 65% → upgraded to Medium"],
], [22, 46, 48, FULL_W - 116], header_fill=DARK, first_col_bold=True)

subhead("Computed is not the same as authorised")
body(
    "The system works out the cadence. It does not act on it until a person signs it off. The Safety "
    "Manager authorises the programme for their site — optionally attaching a specific concern to be "
    "included in scope, which travels all the way to the auditor's brief pack. The Admin then "
    "approves the calendar across all sites."
)
body(
    "Until that authorisation exists, pressing Generate produces nothing and says why. In testing, "
    "all nine sites refused generation with “The programme for this site has not been "
    "authorised”. Generating a year of work nobody signed off is the booking-by-hand problem "
    "wearing a different hat."
)

body("Once authorised, generation produces the year:")
grid([
    ["Site", "Band", "Generated for 2027", "Re-run"],
    ["WindTech Nacelle Manufacturing Ltd", "MEDIUM",
     "5 inspections + 3 audits = 8 events", "0 created, 8 skipped"],
], [58, 20, 56, FULL_W - 134], header_fill=BLUE)

callout(
    "Generation is idempotent, deliberately.",
    "It counts what is already booked in each window and fills only the gap. Running it twice in "
    "March produces the same calendar as running it once — and it never touches an audit that has "
    "already started, because an audit already walked is a record, not a plan.",
    colour=BLUE,
)

screen_note("Web", "Audit Programme — risk band table, per-site authorisation, Generate, "
                   "Admin approval, and the 14-day reminder sweep.")

# ══════════════════════════════════════════════════════════════════════════════
# Step 02
# ══════════════════════════════════════════════════════════════════════════════

part_header("STEP 02", "Somebody has to be independent",
            "PLAN · Safety Manager · HARD STOP")

body(
    "AUD-000010 now exists in the register with no auditor against it, and it cannot move. The "
    "Safety Manager names who audits what — and this is the step the specification marks in red."
)

quote("Names who audits what, and must ensure the auditor is independent of the area being audited. "
      "This is what makes the finding credible.")

body(
    "That is enforced, not advised. Assigning an auditor who is the supervisor of the area under "
    "audit is refused outright:"
)
quote("The auditor cannot be the supervisor of the area being audited — independence is what makes "
      "the finding credible.", source="the system's actual refusal, HTTP 400")

body(
    "The moment a valid auditor is named, two notifications go out. The auditor learns the audit is "
    "theirs and that their brief pack arrives seven days before the visit. The supervisor of the "
    "area — the auditee — is told an audit is coming, with at least two weeks' notice."
)

callout(
    "Unless it is meant to be a surprise.",
    "Two weeks' notice applies to scheduled, management-directed and regulatory audits. It does not "
    "apply to an inspection raised by a risk spike, a score threshold breach or an incident — those "
    "carry no notice by design, and the system stamps nothing rather than recording a notice that "
    "was never given.",
    colour=AMBER, fill="FFFBEB",
)

body("Six things can start an audit, and which one it was changes how it behaves:")
grid([
    ["Trigger", "What it means", "Notice?"],
    ["Scheduled programme", "The annual calendar generated from the site risk band", "Yes"],
    ["After an incident", "A reactive inspection following a significant event", "No"],
    ["Management directed", "Requested by management outside the normal cycle", "Yes"],
    ["Regulatory requirement", "Required by an authority or for certification renewal", "Yes"],
    ["Score threshold breach", "The safety performance score crosses a threshold", "No"],
    ["Risk spike", "Predicted risk rises sharply — raised rather than waiting", "No"],
], [40, FULL_W - 62, 22], header_fill=DARK, first_col_bold=True)

screen_note("Web", "Audit Register — the queue, the step each audit waits on, "
                   "and the assign dialog showing each auditor's workload and qualifications.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# Step 03
# ══════════════════════════════════════════════════════════════════════════════

part_header("STEP 03", "The auditor is briefed before they leave",
            "PREPARE · Auditor · on the phone, readable offline")

body(
    "Seven days before the visit, the system builds the brief pack. The auditor does not assemble "
    "it and cannot edit it. It arrives in their app already containing everything they should walk "
    "in knowing."
)

body("For AUD-000010 the pack contained:")
bullet("previous findings from the last two audits at this site, and whether each was closed")
bullet("open corrective actions still outstanding, flagged where overdue")
bullet("the current score — what this site achieved last time, and its rating")
bullet("25 permits past their validity date at this site, drawn live from the permit register")
bullet("the areas that failed hardest last time, used to reorder the checklist")
bullet("the ISO clauses this audit type maps to")

callout(
    "It is a snapshot, on purpose.",
    "The pack is generated once and stored, rather than recomputed each time it is opened. The "
    "record has to show what the auditor was actually briefed on — not what the data happens to look "
    "like when somebody opens it months later.",
    colour=PURPLE, fill="F5F3FF",
)

subhead("The checklist leads with last time's failures")
body(
    "Sections that produced non-conformances in the previous two audits are weighted, and the "
    "checklist is ordered to put them first. An auditor who runs out of time runs out of it at the "
    "bottom of the list, so what sits at the top matters."
)

subhead("A closed finding is a watchlist item")
body(
    "Findings that were signed off last time appear in a separate repeat watchlist. If the auditor "
    "finds one again, the system flags it as a repeat and treats it as more serious than a first "
    "occurrence — the control was already supposed to be in place."
)

body(
    "The only action on this screen is confirming the brief has been read. That confirmation is what "
    "unlocks the opening meeting: an auditor who has not seen last time's findings cannot set a "
    "credible scope."
)

screen_note("Mobile", "Brief Pack — the whole pack, readable with no signal, ending in "
                      "“I have read the brief”.")

# ══════════════════════════════════════════════════════════════════════════════
# Step 04
# ══════════════════════════════════════════════════════════════════════════════

part_header("STEP 04", "Scope is agreed in the room, not afterwards",
            "CONDUCT · Auditor + Supervisor · on site")

body(
    "The auditor arrives. Before anything is inspected, they hold the opening meeting with the "
    "supervisor of the area, and the app captures it as a structured record rather than a note."
)

body("Three fields are required, and the app refuses to continue without them:")
grid([
    ["Field", "AUD-000010's actual entry"],
    ["Scope", "Assembly hall fire detection, egress and extinguishers. Paint shop out of scope."],
    ["Method", "Physical walk, worker interview, record sampling"],
    ["Sampling approach", "All 12 extinguishers; 3 of 9 exit routes at random; 5 worker interviews"],
], [34, FULL_W - 34], header_fill=BLUE, first_col_bold=True)

body(
    "They are three separate required fields rather than one notes box because the entire value of "
    "the opening meeting is that there is no dispute afterwards about what was in or out of scope — "
    "and free text does not settle that argument."
)

body(
    "The supervisor attends, and their attendance is recorded either way. If they are not there the "
    "audit is not blocked, but the fact is stored: “scope was agreed jointly” is a claim "
    "the report makes, and it has to be true. The meeting is GPS-stamped."
)

screen_note("Mobile", "Opening Meeting — scope, method, sampling, attendees, and whether the "
                      "supervisor was present.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# Steps 05-06
# ══════════════════════════════════════════════════════════════════════════════

part_header("STEPS 05–06", "The walk",
            "CONDUCT · Auditor · one item per screen, offline-capable")

body(
    "Now the auditor walks the site. The app shows one checklist item per screen and they swipe "
    "through. Each item carries its section, its ISO clause, and — where it applies — a red flag "
    "marking it as critical."
)

body("AUD-000010's fire safety checklist had four items. Here is what happened, in order:")

grid([
    ["#", "Item", "Answer", "Running score"],
    ["1", "Extinguishers  (CRITICAL)", "Full — 2 pts", "100%  Excellent"],
    ["2", "Exit Routes  (CRITICAL)", "None — 0 pts", "50%  Poor"],
    ["3", "Alarm System", "Partial — 1 pt", "50%  Poor"],
    ["4", "Emergency Lighting", "Full — 2 pts", "62.5%  Acceptable"],
], [10, 62, 34, FULL_W - 106], header_fill=DARK)

subhead("What happened when item 2 scored zero")
body(
    "Exit Route 3 was blocked by stacked pallets. The auditor tapped None. Because that item is "
    "marked critical, three things happened before their thumb left the screen:"
)
bullet("the audit moved into stop-work, which the eight-stage engine reads as RESPOND")
bullet("the Safety Manager and the executive were notified immediately — not at submit, not in the report")
bullet("the app showed the auditor exactly what had been raised and by what rule")

quote("'Exit Routes' is a critical item and scored zero. The Safety Manager and the executive have "
      "been notified. Work may be suspended before this audit finishes.",
      source="what the app actually displayed")

body(
    "The auditor confirms the hazard is contained and the walk resumes — it does not restart. "
    "Containment happens first; the audit carries on afterwards."
)

subhead("The scoring rubric")
body(
    "Answers are not pass/fail. A full conformance is worth twice an observation, and items marked "
    "Not Applicable leave the calculation entirely so a score is never diluted by questions that did "
    "not apply to that site."
)

grid([
    ["Answer", "Points", "Meaning"],
    ["Full compliance", "2", "Meets the requirement"],
    ["Partial / observation", "1", "Compliant but improvable"],
    ["Non-compliance", "0", "Does not meet the requirement"],
    ["Not applicable", "excluded", "Removed from the denominator entirely"],
], [38, 22, FULL_W - 60], header_fill=BLUE, first_col_bold=True)

body(
    "Both specifications express the formula differently and they are mathematically identical. The "
    "app ships both so an auditor holding either document recognises what the phone is showing:"
)
grid([
    ["Specification", "Formula"],
    ["Algorithmic spec (ALGO-ISMS-WF-2026)", "(points earned / points possible) x 100"],
    ["Form specification (AUD-FORM-01)", "(C + 0.5 x OBS) / total assessed x 100"],
], [70, FULL_W - 70], header_fill=DARK, first_col_bold=True)

subhead("Evidence is attached to the line it proves")
quote("Photos attached to the specific checklist item, not dumped in a general folder.")

body("On AUD-000010 the auditor attached three pieces of evidence, each bound to its checklist line:")
grid([
    ["Kind", "Attached to", "What it recorded"],
    ["Note", "Exit Routes", "Pallets stacked across exit 3"],
    ["Scan", "Alarm System", "Asset reference AST-EXT-0442"],
    ["Interview", "Exit Routes",
     "Worker named exit 3 — the blocked one. Did not know the alternate. Competence card not verified."],
], [22, 34, FULL_W - 56], header_fill=BLUE, first_col_bold=True)

callout(
    "A worker interview is evidence, not a note.",
    "The specification is blunt about this: what the worker actually does is the evidence, not what "
    "the procedure says. So the answer is recorded against the person who gave it, against the "
    "checklist line it proves, with the prompt they were asked and whether their competence card "
    "checked out against the matrix.",
    colour=ORANGE, fill=LIGHT_BG,
)

body(
    "Every answer and every piece of evidence is GPS-stamped, which is what makes the observation "
    "defensible to a regulator: it proves the observation was made at the place claimed."
)

callout(
    "The whole walk works with no signal.",
    "Audits happen in plant areas, tank farms and remote sites. Answers are held on the device the "
    "instant they are tapped. Where there is signal they also go up immediately — that is what "
    "returns the running score and fires the critical alert while work can still be suspended. Where "
    "there is not, the answer is marked unsent and the whole set is carried up at step 7.",
    colour=GREEN, fill="ECFDF5",
)

screen_note("Mobile", "Field Inspection — one item per screen, swipe navigation, running score in "
                      "the header, evidence sheet, and the critical-item alert.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# Step 07
# ══════════════════════════════════════════════════════════════════════════════

part_header("STEP 07", "The auditor judges; the system counts",
            "CLASSIFY · Auditor · HARD STOP")

body(
    "With the walk finished, the auditor reviews every answer and assigns a classification to each "
    "finding. The app suggests one; the auditor decides."
)

grid([
    ["Classification", "What it means", "Action due"],
    ["CONFORMANCE", "Meets the requirement. Recorded as a positive — audits record what is working, "
     "not only what is wrong.", "—"],
    ["OBSERVATION", "Compliant but improvable. A finding, but not a non-conformance.", "—"],
    ["MINOR NC", "A lapse that does not undermine the system.", "30 days"],
    ["MAJOR NC", "A systemic failure. Safety Manager notified within 24 hours.", "7 days"],
    ["CRITICAL / REGULATORY", "Immediate danger or a legal breach. Executive notified at once and "
     "work may be suspended.", "1 day"],
], [40, FULL_W - 62, 22], header_fill=DARK, first_col_bold=True)

subhead("Two judgements are not the auditor's to make")
body(
    "The auditor owns every classification except where a rule overrides them, and the app says so "
    "rather than silently correcting:"
)
bullet("A critical item scoring zero is a Major NC at minimum. It can be escalated to Critical; it "
       "cannot be softened below Major. That is the entire point of marking an item critical.",
       bold_lead="Critical zeros. ")
bullet("A section scoring below 60% raises a Minor NC of its own, attributed to nobody and "
       "impossible to suppress. A section falling below the threshold is a lapse in the system, not "
       "in one item.", bold_lead="Weak sections. ")

body("On AUD-000010, both rules fired. The Egress section scored 50%:")

grid([
    ["Section", "Score", "Outcome"],
    ["Detection & Suppression", "75%", "—"],
    ["Egress", "50%", "Below the 60% threshold → automatic Minor NC"],
], [46, 20, FULL_W - 66], header_fill=BLUE, first_col_bold=True)

body("The five findings the system produced:")
grid([
    ["Ref", "Finding", "Classification", "Action due"],
    ["AUD-000010-F01", "Extinguishers", "Conformance", "—"],
    ["AUD-000010-F02", "Exit Routes", "MAJOR NC  (critical item, auto)", "28 Aug — 7 days"],
    ["AUD-000010-F03", "Alarm System", "Observation", "—"],
    ["AUD-000010-F04", "Emergency Lighting", "Conformance", "—"],
    ["AUD-000010-F05", "Section 'Egress' scored 50%", "MINOR NC  (section rule, auto)", "20 Sep — 30 days"],
], [34, 46, 50, FULL_W - 130], header_fill=DARK)

subhead("The score and the rating are different things")
body(
    "The system computed 62.5% — five points earned out of eight possible — which is the Acceptable "
    "band. The overall rating came back Unsatisfactory."
)

callout(
    "62% and Unsatisfactory at the same time.",
    "That is not a contradiction; it is the design. The rating is set from the finding counts, not "
    "from the percentage. Any Major non-conformance or regulatory breach makes an audit "
    "unsatisfactory whatever it scored — because a good average with a systemic failure inside it is "
    "still a site with a systemic failure. More than three Minor NCs makes it Requires Improvement.",
    colour=RED, fill="FEF2F2",
)

grid([
    ["Band", "Range", "AUD-000010"],
    ["EXCELLENT", "90% and above", ""],
    ["GOOD", "75 – 89%", ""],
    ["ACCEPTABLE", "60 – 74%", "62.5%  ←"],
    ["POOR", "Below 60%", ""],
], [28, 40, FULL_W - 68], header_fill=DARK, first_col_bold=True)

body(
    "Two thresholds sit alongside the bands and act on their own: below 70% automatically alerts the "
    "Safety Manager, and below 65% twice in a row at one site forces a re-audit. At 62%, the first "
    "of those fired."
)

screen_note("Mobile", "Findings & Score — the arithmetic shown but not editable, the "
                      "classification chips that are, and the two rules the auditor cannot override.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# Step 08
# ══════════════════════════════════════════════════════════════════════════════

part_header("STEP 08", "The auditee gets to say it is wrong",
            "AGREE · Auditor + Supervisor · HARD STOP · before anyone leaves site")

body(
    "The closing meeting is where the supervisor sees the findings — on the auditor's screen, in "
    "front of them — and gets their one opportunity to correct a factual error before anything is "
    "fixed or reported."
)

quote("Their opportunity to correct a factual error before anything is fixed. After this meeting "
      "findings are locked and can only change through a formal amendment.")

body("Three things have to happen, and the app will not finish without all three:")
bullet("every finding is presented on screen to the supervisor", bold_lead="Presented. ")
bullet("the supervisor confirms the findings are factually accurate — or disputes them",
       bold_lead="Confirmed. ")
bullet("both parties sign on the device, and the findings lock immediately",
       bold_lead="Signed. ")

subhead("Disputing is a real outcome")
body(
    "If the supervisor says a finding is factually wrong, the dispute is recorded and nothing locks. "
    "The auditor goes back, corrects the error and holds the meeting again. That was tested on "
    "AUD-000010: a disputed closing meeting left findings unlocked and the audit sitting at step 8, "
    "exactly as it should."
)

subhead("Both signatures, on the device, before anyone leaves")
body(
    "Signing later — on paper, or by email the next day — is the gap through which findings drift "
    "between the walk and the report. So the signature pad is on the phone, and capturing both "
    "signatures is what locks the findings."
)

body(
    "The corrective action deadlines are also agreed here, in front of the supervisor, and they "
    "override the system's defaults. The supervisor committed to those dates with the auditor "
    "standing there."
)

body("After the meeting, the checklist is frozen. Attempting to change an answer returns:")
quote("Findings were locked at the closing meeting and can only change through a formal amendment.",
      source="the system's actual refusal, HTTP 400")

screen_note("Mobile", "Closing Meeting — findings list, agreed timeframes, the accuracy decision, "
                      "and two on-device signature pads.")

# ══════════════════════════════════════════════════════════════════════════════
# Step 09
# ══════════════════════════════════════════════════════════════════════════════

part_header("STEP 09", "One signature does four things",
            "REPORT · Auditor · HARD STOP")

quote("The report cannot be issued without the auditor's signature. Signing triggers distribution "
      "and creates the corrective actions.")

body(
    "Signing is not a formality on top of a transition — it is the transition. When the auditor "
    "signed AUD-000010, four things happened in one movement:"
)
bullet("the report was generated from the data — scores, findings, benchmark, clause mapping")
bullet("a corrective action was created for every non-conformance, with the deadlines agreed at the "
       "closing meeting")
bullet("the report was distributed to the supervisor and the Safety Manager")
bullet("the site's history was re-evaluated for the persistent-poor-performance rule")

body("The report the system built, with no document written by hand:")
grid([
    ["Element", "AUD-000010"],
    ["Reference", "RPT-000010"],
    ["Score", "62.5% — Acceptable"],
    ["Overall rating", "Unsatisfactory"],
    ["Findings", "5 (2 conformance · 1 observation · 1 minor NC · 1 major NC)"],
    ["Benchmark", "−20.5 points against AUD-000009's 83%"],
    ["Repeat findings", "0"],
    ["Clause mapping", "ISO 45001 8.2 — 3 findings, worst Major NC · OSHA 1910.157 — 1 finding"],
    ["Corrective actions", "CAPA-000236 (due 26 Aug) · CAPA-000237 (due 20 Sep)"],
    ["Signed by", "Auditor One · factual accuracy confirmed by Supervisor One"],
], [34, FULL_W - 34], header_fill=BLUE, first_col_bold=True)

callout(
    "There is no way to have one without the other.",
    "No path issues a report without raising the corrective actions, and none raises actions from a "
    "report nobody signed. Tying them to the same act is what stops findings being reported and then "
    "quietly forgotten.",
    colour=BLUE,
)

subhead("Then it moves to the desktop")
body(
    "The full report — a long document with sections, findings, benchmark comparison and clause "
    "mapping — is reviewed and distributed from the web console, where a long document is genuinely "
    "easier to work with. Two gates apply, in order:"
)
grid([
    ["Who", "What they do", "Gate"],
    ["Safety Manager", "Reviews and approves the report before wider distribution",
     "Cannot happen before the report is signed"],
    ["Admin", "Owns distribution beyond the site", "Refused until the Safety Manager has approved"],
], [30, 62, FULL_W - 92], header_fill=DARK, first_col_bold=True)

body(
    "That second gate was tested: distributing before approval returns “The Safety Manager "
    "reviews and approves the report before wider distribution”. Without it, the review would "
    "be decorative."
)

screen_note("Mobile", "Audit Report — the summary, the signature pad, and Sign & Issue.")
screen_note("Web", "Report Review — the full document, approve for distribution, release "
                   "beyond the site, print.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# Step 10
# ══════════════════════════════════════════════════════════════════════════════

part_header("STEP 10", "The audit stays open",
            "CLOSE · The System + Auditor · HARD STOP")

body(
    "AUD-000010's report was issued. The audit did not close. It moved to step 10 with two open "
    "findings and stayed there."
)

quote("Keeps the audit open until every action is verified. Confirms findings were genuinely "
      "closed, not just marked closed.")

subhead("Marking complete is a claim; verification is a visit")
body(
    "The auditor returns at 30, 60 and 90 days and checks each fix on site. The system will not let "
    "them verify a finding whose corrective action is still open:"
)
quote("Corrective action CAPA-000236 is still Open. It has to be completed before the finding can "
      "be verified.", source="the system's actual refusal, HTTP 400")

body(
    "Once both actions were completed and both findings verified effective, the audit closed — "
    "status completed, all ten steps done, stage 8 CLOSE."
)

subhead("And if the fix is not holding")
body(
    "A failed effectiveness check does not just reopen the finding. It reopens the corrective action, "
    "and it reopens the audit itself — even one already closed. Tested on AUD-000010: a failed "
    "60-day check took a completed audit back to capa_open, step 10 blocked, stage 5 IMPROVE."
)

callout(
    "This was a real bug, found and fixed.",
    "The first implementation treated an audit as closed if its status said completed, which meant "
    "clearing the closure date did nothing — the status the system was about to recompute answered "
    "“still closed” and overwrote the reopen with itself. A migration backfilled closure "
    "dates onto historic audits so a single field is the authority.",
    colour=AMBER, fill="FFFBEB",
)

screen_note("Mobile", "Track Findings Out — per-finding verification with GPS, and the "
                      "cross-audit 30/60/90 queue.")
screen_note("Web", "Trends & Oversight — cross-site comparison, repeat-finding analysis, "
                   "and the re-audit decision.")

# ══════════════════════════════════════════════════════════════════════════════
# Escalations
# ══════════════════════════════════════════════════════════════════════════════

part_header("", "The five rules that fire on their own",
            "Each fires without anyone needing to notice.")

body(
    "None of these is a screen somebody has to open or a report somebody has to run. Each is raised "
    "by the transition that creates the condition, at the moment it is created."
)

grid([
    ["Trigger", "What happens", "Fires at"],
    ["Critical finding on site",
     "Safety Manager and executive notified immediately. Work may be suspended before the audit finishes.",
     "The moment the item is answered"],
    ["Major non-conformance",
     "Safety Manager notified within 24 hours. A corrective action must exist within 7 days.",
     "Classification, step 7"],
    ["Regulatory finding", "Triggers the statutory notification workflow with its own legal deadline.",
     "Classification, step 7"],
    ["Audit not conducted",
     "Alert at 110% of the scheduled date. A missed audit is itself a finding.",
     "Daily sweep"],
    ["Persistent poor performance",
     "Two Major NCs at one site in 12 months, or below 65% twice running → mandatory re-audit within 30 days.",
     "Report issue, step 9"],
], [38, FULL_W - 74, 36], header_fill=RED, first_col_bold=True)

subhead("The re-audit decision belongs to a person")
body(
    "The trigger fires automatically. What to do about it does not. The Safety Manager either "
    "schedules the re-audit — which creates it, seeds its checklist and links it to the original — "
    "or waives it. Waiving requires a written reason, and the system refuses without one."
)
body(
    "An unexplained waiver of a mandatory re-audit is the single most useful thing for a regulator "
    "to find, which is precisely why it cannot be recorded silently.",
    italic=True,
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# Roles
# ══════════════════════════════════════════════════════════════════════════════

part_header("", "Who does what", "The specification's job titles, mapped to this platform's roles.")

body(
    "The source documents name several audit roles — Lead Auditor, Audit Team, Auditee Management, "
    "Safety Advisor, ISMS Director. Those are job titles in a large safety department, not app "
    "roles. On this platform they resolve to the four mobile roles plus web admin."
)

grid([
    ["Document role", "Platform role", "Note"],
    ["Lead Auditor + Audit Team", "AUDITOR (mobile)",
     "One role; where several are assigned to one audit, one is designated lead"],
    ["Auditee Management", "SUPERVISOR (mobile)", "The supervisor of the area under audit"],
    ["Workers interviewed", "WORKER (mobile)", "Observed and questioned during the walk"],
    ["Safety Manager / Safety Advisor", "SAFETY MANAGER (mobile + web)", "Authorises, assigns, approves"],
    ["ISMS Director / Safety Director", "ADMIN (web only)",
     "No mobile app; oversight and programme ownership only"],
], [48, 44, FULL_W - 92], header_fill=DARK, first_col_bold=True)

subhead("In one line each")
grid([
    ["AUDITOR", "Conducts the field inspection · collects and links evidence · classifies every "
                "finding · signs the report — nothing issues without it · verifies effectiveness at 30/60/90"],
    ["WORKER", "Interviewed and observed · shows competence and records · completes actions assigned "
               "to them · receives the toolbox talk"],
    ["SUPERVISOR", "Gets two weeks' notice · attends opening and closing · confirms factual accuracy "
                   "· owns the actions for their area"],
    ["SAFETY MANAGER", "Authorises the programme · assigns an independent auditor · approves the "
                       "report · owns the re-audit decision"],
    ["ADMIN", "Owns the audit programme · maintains checklist templates · receives every report · "
              "owns cross-site trends"],
    ["THE SYSTEM", "Builds the schedule and brief · calculates scores and thresholds · raises actions "
                   "from findings · fires the re-audit trigger"],
], [32, FULL_W - 32], header_fill=ORANGE, first_col_bold=True)

# ══════════════════════════════════════════════════════════════════════════════
# Where it runs
# ══════════════════════════════════════════════════════════════════════════════

part_header("", "Why the work splits where it does")

quote("Everything requiring observation, evidence or a signature happens where the work is. "
      "Everything requiring reading, comparison or distribution happens where the screen is bigger. "
      "The auditor never re-enters anything they already recorded in the field.")

grid([
    ["Step", "Where", "Screen"],
    ["01  Schedule generated", "Web", "Audit Programme"],
    ["02  Team assigned", "Web", "Audit Register"],
    ["03  Pre-audit prep", "Mobile", "Brief Pack"],
    ["04  Opening meeting", "Mobile", "Opening Meeting"],
    ["05  Field inspection", "Mobile", "Field Inspection"],
    ["06  Evidence captured", "Mobile", "Field Inspection · Worker Interview"],
    ["07  Findings & score", "Mobile", "Findings & Score"],
    ["08  Closing meeting", "Mobile", "Closing Meeting"],
    ["09  Report issued", "Mobile signs · Web reviews", "Audit Report · Report Review"],
    ["10  Findings tracked out", "Mobile verifies · Web oversees", "Track Findings Out · Trends"],
], [44, 46, FULL_W - 90], header_fill=DARK, first_col_bold=True)

subhead("What the phone has to be able to do")
grid([
    ["Capability", "Why", "Status"],
    ["Offline execution", "Audits happen where there is no signal. Not optional for a field tool.", "Built"],
    ["Camera & evidence linking", "Photos attached to the checklist item, not a folder.", "Built"],
    ["GPS stamping", "Proves the observation was made where claimed.", "Built"],
    ["On-device signature", "Both parties sign before leaving site, so findings lock immediately.", "Built"],
    ["Voice-to-text", "Hands are often gloved or occupied.", "OS dictation"],
    ["QR / barcode scan", "Identify equipment without typing a reference.", "Typed entry — needs a scanner library"],
], [40, FULL_W - 88, 48], header_fill=GREEN, first_col_bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# What was built
# ══════════════════════════════════════════════════════════════════════════════

part_header("", "What was built", "For the record — the shape of the implementation.")

subhead("Data")
grid([
    ["Table", "Holds"],
    ["audits", "The audit and its whole lifecycle — trigger, band, meetings, signatures, score, report"],
    ["audit_checklist_items", "The checklist as rows, so a section can score on its own"],
    ["audit_findings", "One row per classified finding, tracked out individually"],
    ["audit_evidence", "Photos, documents, scans, notes and worker interviews, bound to a checklist line"],
    ["audit_programme", "Each site's band, cadence, authorisation and generation history"],
    ["audit_checklist_templates", "The versioned templates every audit runs from"],
], [52, FULL_W - 52], header_fill=DARK, first_col_bold=True)

subhead("Judgement, kept out of the controllers")
grid([
    ["Module", "Owns"],
    ["audit_scoring", "The rubric, the bands, the five classifications, the overall rating"],
    ["audit_programme", "The six triggers, the four bands, the frequency table"],
    ["audit_steps", "The ten steps, their owners, their hard stops, and the status they imply"],
    ["audit_escalation", "The five triggers that fire on their own"],
    ["audit_brief", "The step 03 brief pack"],
    ["audit_calendar", "Calendar generation and the 14-day reminder"],
    ["audit_trends", "Cross-site comparison and repeat-finding analysis"],
    ["audit_templates", "Template resolution and versioning"],
], [40, FULL_W - 40], header_fill=BLUE, first_col_bold=True)

subhead("Screens")
grid([
    ["Mobile (Auditor)", "Web console"],
    ["Dashboard — queue, programme, escalations", "Audit Programme — bands, authorisation, generation"],
    ["Assigned Audits — step state per audit", "Audit Register — the queue and assignment"],
    ["Audit Detail — the one action that is due", "Report Review — the long document, approve, distribute"],
    ["Brief Pack · Opening Meeting", "Trends & Oversight — cross-site, repeats, re-audit"],
    ["Field Inspection · Worker Interview", "Templates & Auditors — versioned editor, register"],
    ["Findings & Score · Closing Meeting", ""],
    ["Audit Report · Track Findings Out", ""],
], [FULL_W // 2, FULL_W // 2], header_fill=ORANGE)

subhead("Honest limits")
body("Three things are worth stating plainly rather than leaving to be discovered:")
bullet("QR and barcode scanning is typed entry. No scanner library is installed, and shipping a "
       "button that pretends to scan would be worse than one that plainly asks for the reference.",
       bold_lead="No scanner. ")
bullet("Voice-to-text relies on the phone keyboard's dictation key rather than an in-app speech "
       "integration.", bold_lead="Dictation, not integration. ")
bullet("The mobile app is type-checked and the web app builds, but neither has been walked through "
       "on a device against live data. The backend flow has — every number in this document came "
       "from that run.", bold_lead="Not device-tested. ")

spacer(10)

t = doc.add_table(rows=1, cols=1)
c = t.rows[0].cells[0]
c.text = ""
p = c.paragraphs[0]
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)
r = p.add_run(
    "The auditor is one of four mobile roles. Steps 4 to 8 are conducted in the field on the phone, "
    "offline-capable, with evidence and signatures captured on site. An audit is not closed when the "
    "report is issued — it stays open until every corrective action it raised has been verified "
    "effective."
)
r.italic = True
r.font.size = Pt(9.5)
r.font.color.rgb = RGBColor.from_string(WHITE)
shade(c, DARK)
set_widths(t, [FULL_W])

out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "HSE_Auditor_Flow_Story.docx")
doc.save(out)
print(f"Written: {out}")
